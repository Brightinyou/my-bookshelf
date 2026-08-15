"""위키 노트(마크다운) → Word(.docx) 내보내기 (Phase 2, 2026-07-24).

옵시디언을 쓰지 않는 사용자를 위해, 허브 노트와 동일한 내용을 편집 가능한
Word 문서로 저장한다. python-docx만 사용(순수 파이썬·시스템 의존성 없음)."""
from __future__ import annotations

import json
import re
import tempfile
import shutil
from pathlib import Path

import config as cfg


def set_docx_dir(path_str: str) -> None:
    """~/.config/mybookshelf/config.json의 dirs.docx 갱신 — 앱 재시작 후 적용.
    (옵시디언 보관함 설정의 set_wiki_dir와 동일한 방식, 2026-07-25)"""
    f = cfg.CONFIG_FILE
    try:
        d = json.loads(f.read_text(encoding="utf-8")) if f.exists() else {}
    except Exception:
        d = {}
    d.setdefault("dirs", {})["docx"] = path_str
    f.parent.mkdir(parents=True, exist_ok=True)
    f.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")


def _safe_name(stem: str) -> str:
    return re.sub(r'[/\\:*?"<>|]', "_", stem).strip() or "wiki"


_EA_FONT = "맑은 고딕"


def _set_ea_font(run, name: str = _EA_FONT) -> None:
    """런의 동아시아(w:eastAsia) 글꼴을 명시적으로 지정.

    python-docx 기본 템플릿은 Normal/제목/인용 등 대부분 스타일에서 동아시아 글꼴을
    테마의 빈 값(<a:ea typeface="">)에 맡긴다. MS Word는 이를 한글 폰트로 암묵
    대체해주지만 다른 뷰어·언어팩 환경에서는 대체되지 않아 한글이 네모로 깨진다
    (표 셀만 다른 경로로 우연히 정상 렌더링됨). 그래서 런마다 직접 지정한다."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_font_fallback_chain(docx_path: Path) -> None:
    """fontTable.xml에 맑은 고딕→굴림→바탕 대체(altName) 체인을 등록.

    rFonts는 폰트 하나만 지정할 수 있어, 지정한 동아시아 폰트가 없는 환경에서도
    Word가 자동으로 다음 폰트를 시도하도록 OOXML altName 메커니즘을 별도로 쓴다."""
    import zipfile
    from lxml import etree
    from docx.oxml.ns import qn as _qn

    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    root = etree.fromstring(data["word/fontTable.xml"])
    existing = {f.get(_qn("w:name")) for f in root.findall(_qn("w:font"))}

    def _ensure(name: str, alt: str | None) -> None:
        if name in existing:
            return
        el = etree.SubElement(root, _qn("w:font"))
        el.set(_qn("w:name"), name)
        if alt:
            alt_el = etree.SubElement(el, _qn("w:altName"))
            alt_el.set(_qn("w:val"), alt)

    _ensure(_EA_FONT, "굴림")
    _ensure("굴림", "바탕")

    data["word/fontTable.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])


def _add_inline(paragraph, text: str) -> None:
    """**굵게** 정도만 처리한 인라인 런 추가."""
    for i, seg in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = paragraph.add_run(seg[2:-2])
            r.bold = True
        else:
            r = paragraph.add_run(seg)
        _set_ea_font(r)


def note_md_to_docx(md: str, out_path: Path, *, meta: dict | None = None) -> Path:
    """마크다운 노트 문자열을 .docx로 변환해 out_path에 저장."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # ── frontmatter 분리 → 제목·서지 헤더 ──
    body = md
    fm: dict[str, str] = {}
    m = re.match(r"^---\n(.*?)\n---\n?(.*)$", md, re.S)
    if m:
        for line in m.group(1).splitlines():
            mm = re.match(r"^(\w+):\s*(.*)$", line)
            if mm:
                fm[mm.group(1)] = mm.group(2).strip().strip('"')
        body = m.group(2)
    fm.update(meta or {})

    title = fm.get("title") or (meta or {}).get("title") or ""
    if title:
        _h = doc.add_heading(title, level=0)
        for _r in _h.runs:
            _set_ea_font(_r)
    _bib = " · ".join(
        x for x in [fm.get("author", ""), fm.get("published", ""), fm.get("publisher", "")] if x
    )
    if _bib:
        _p = doc.add_paragraph()
        _r = _p.add_run(_bib)
        _r.italic = True
        _r.font.size = Pt(10)
        _set_ea_font(_r)

    # ── 본문 라인 단위 변환 ──
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1
            continue
        # 표(| ... |) — 연속 파이프 줄 묶기 (구분선 |---| 은 건너뜀)
        if ln.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^\s*:?-{2,}", cells[0] if cells else ""):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncol)
                tbl.style = "Light Grid Accent 1"
                for r in rows:
                    cells = tbl.add_row().cells
                    for c in range(ncol):
                        cells[c].text = r[c] if c < len(r) else ""
                        for _p in cells[c].paragraphs:
                            for _r in _p.runs:
                                _set_ea_font(_r)
            continue
        # 헤딩
        h = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if h:
            level = min(len(h.group(1)), 4)
            _h = doc.add_heading(h.group(2), level=level)
            for _r in _h.runs:
                _set_ea_font(_r)
            i += 1
            continue
        # 인용
        if ln.lstrip().startswith(">"):
            _p = doc.add_paragraph(style="Intense Quote")
            _add_inline(_p, ln.lstrip()[1:].strip())
            i += 1
            continue
        # 키워드 해시태그 줄 (#키워드 — 해설)
        if ln.lstrip().startswith("#") and not ln.lstrip().startswith("##"):
            _p = doc.add_paragraph()
            _r = _p.add_run(ln.strip())
            _r.bold = True
            _set_ea_font(_r)
            i += 1
            continue
        # 불릿
        if re.match(r"^\s*[-*]\s+", ln):
            _p = doc.add_paragraph(style="List Bullet")
            _add_inline(_p, re.sub(r"^\s*[-*]\s+", "", ln))
            i += 1
            continue
        # 일반 문단
        _p = doc.add_paragraph()
        _add_inline(_p, ln)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _add_font_fallback_chain(out_path)
    return out_path


def build_docx_from_chapter_summaries(ws_name: str, stem: str, out_dir: Path) -> tuple[bool, str]:
    """챕터 요약 → 허브 노트(임시 생성) → .docx. (ok, path or msg)."""
    from services.wiki import build_wiki_from_chapter_summaries
    tmp = Path(tempfile.mkdtemp(prefix="mb_docx_"))
    try:
        ok, msg = build_wiki_from_chapter_summaries(ws_name, stem, wiki_dir=tmp)
        if not ok:
            return False, msg
        md_path = Path(msg)
        if not md_path.exists():
            return False, "노트 생성 실패"
        md = md_path.read_text(encoding="utf-8")
        out_path = out_dir / (_safe_name(stem) + ".docx")
        note_md_to_docx(md, out_path)
        return True, str(out_path)
    except Exception as e:
        return False, f"{type(e).__name__}: {str(e)[:150]}"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
